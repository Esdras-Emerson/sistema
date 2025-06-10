import streamlit as st
import os
from io import BytesIO
from pathlib import Path
import tempfile
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores.faiss import FAISS
from langchain_openai.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
import pandas as pd
from langchain.prompts import PromptTemplate

def chat_interface():
    st.header('💬 Converse Comigo - Upload e Consulta de Arquivos')
    
    # Verificar OpenAI API key
    openai_api_key = st.secrets.get("OPENAI_API_KEY", "")
    if not openai_api_key:
        st.error("❌ Chave da API OpenAI não configurada. Adicione OPENAI_API_KEY no secrets.toml")
        return
    
    os.environ["OPENAI_API_KEY"] = openai_api_key
    
    st.info("📁 Faça upload de arquivos para consultar informações específicas deles")
    
    # Upload múltiplo de arquivos
    uploaded_files = st.file_uploader(
        'Selecione arquivos para consulta (Word, PDF, Excel)',
        type=['docx', 'doc', 'pdf', 'xls', 'xlsx'],
        accept_multiple_files=True,
        help="Você pode selecionar múltiplos arquivos de diferentes tipos"
    )
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} arquivo(s) selecionado(s)")
        
        # Mostrar arquivos selecionados
        with st.expander("📋 Arquivos selecionados"):
            for arquivo in uploaded_files:
                st.write(f"• {arquivo.name} ({arquivo.size} bytes)")
        
        # Botão para processar arquivos
        if st.button("🔄 Processar Arquivos e Inicializar Chat", type="primary"):
            processar_e_inicializar_chat(uploaded_files)
    
    # Interface de chat (só aparece se já tiver arquivos processados)
    if "chat_chain_upload" in st.session_state:
        st.divider()
        st.subheader("💭 Chat - Consulte os arquivos carregados")
        
        # Mostrar informações dos arquivos processados
        if "arquivos_processados" in st.session_state:
            with st.expander("📊 Status dos arquivos processados"):
                for info in st.session_state.arquivos_processados:
                    st.write(f"✅ {info['nome']} - {info['tipo']} - {info['status']}")
        
        # Histórico de mensagens
        if "chat_history_upload" not in st.session_state:
            st.session_state.chat_history_upload = []
        
        # Mostrar mensagens anteriores
        for message in st.session_state.chat_history_upload:
            with st.chat_message(message["role"]):
                st.write(message["content"])
                if message["role"] == "assistant" and "sources" in message:
                    with st.expander("📄 Fontes consultadas"):
                        for source in message["sources"]:
                            st.write(f"• {source}")
        
        # Input do usuário
        user_input = st.chat_input("Digite sua pergunta sobre os arquivos...")
        
        if user_input:
            # Adicionar pergunta do usuário
            st.session_state.chat_history_upload.append({"role": "user", "content": user_input})
            
            with st.chat_message("user"):
                st.write(user_input)
            
            # Processar resposta
            with st.chat_message("assistant"):
                with st.spinner("🔍 Consultando arquivos..."):
                    try:
                        response = st.session_state.chat_chain_upload({"query": user_input})
                        answer = response["result"]
                        sources = [doc.metadata.get("source", "Desconhecido") for doc in response.get("source_documents", [])]
                        
                        st.write(answer)
                        
                        if sources:
                            with st.expander("📄 Fontes consultadas"):
                                for source in set(sources):
                                    st.write(f"• {source}")
                        
                        # Adicionar resposta ao histórico
                        st.session_state.chat_history_upload.append({
                            "role": "assistant", 
                            "content": answer,
                            "sources": list(set(sources))
                        })
                        
                    except Exception as e:
                        st.error(f"Erro ao processar pergunta: {str(e)}")
        
        # Botões de controle
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Limpar Chat"):
                st.session_state.chat_history_upload = []
                st.rerun()
        
        with col2:
            if st.button("🔄 Carregar Novos Arquivos"):
                # Limpar sessão para permitir novos uploads
                for key in ["chat_chain_upload", "arquivos_processados", "chat_history_upload"]:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()
    
    else:
        if not uploaded_files:
            st.info("👆 Selecione arquivos acima para começar")
        else:
            st.info("👆 Clique em 'Processar Arquivos' para inicializar o chat")

def processar_e_inicializar_chat(uploaded_files):
    """Processa os arquivos carregados e inicializa o chat"""
    
    with st.spinner("📄 Processando arquivos..."):
        documentos = []
        arquivos_processados = []
        
        progress_bar = st.progress(0)
        
        for i, arquivo in enumerate(uploaded_files):
            st.write(f"Processando: {arquivo.name}")
            
            try:
                # Processar baseado no tipo de arquivo
                if arquivo.name.lower().endswith('.pdf'):
                    docs = processar_pdf(arquivo)
                    tipo = "PDF"
                elif arquivo.name.lower().endswith(('.docx', '.doc')):
                    docs = processar_word(arquivo)
                    tipo = "Word"
                elif arquivo.name.lower().endswith(('.xlsx', '.xls')):
                    docs = processar_excel(arquivo)
                    tipo = "Excel"
                else:
                    st.warning(f"Tipo de arquivo não suportado: {arquivo.name}")
                    continue
                
                if docs:
                    documentos.extend(docs)
                    arquivos_processados.append({
                        "nome": arquivo.name,
                        "tipo": tipo,
                        "status": f"{len(docs)} documento(s) extraído(s)"
                    })
                    st.success(f"✅ {arquivo.name} processado")
                else:
                    st.error(f"❌ Falha ao processar {arquivo.name}")
                    arquivos_processados.append({
                        "nome": arquivo.name,
                        "tipo": tipo,
                        "status": "Erro no processamento"
                    })
                
            except Exception as e:
                st.error(f"❌ Erro ao processar {arquivo.name}: {str(e)}")
                arquivos_processados.append({
                    "nome": arquivo.name,
                    "tipo": "Erro",
                    "status": f"Erro: {str(e)}"
                })
            
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        if documentos:
            # Criar vector store e chat
            criar_chat_com_documentos(documentos)
            st.session_state.arquivos_processados = arquivos_processados
            st.success(f"🎉 Chat inicializado com {len(documentos)} documento(s)!")
            st.rerun()
        else:
            st.error("❌ Nenhum documento foi processado com sucesso")

def processar_pdf(arquivo):
    """Processa arquivo PDF"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(arquivo.read())
            tmp_file_path = tmp_file.name
        
        loader = PyPDFLoader(tmp_file_path)
        docs = loader.load()
        
        # Adicionar metadata
        for doc in docs:
            doc.metadata['source'] = arquivo.name
            doc.metadata['tipo'] = 'PDF'
        
        os.unlink(tmp_file_path)
        return docs
        
    except Exception as e:
        st.error(f"Erro ao processar PDF {arquivo.name}: {str(e)}")
        return []

def processar_word(arquivo):
    """Processa arquivo Word"""
    try:
        # Tentar diferentes métodos para Word
        try:
            from docx import Document as DocxDocument
            
            # Método 1: python-docx
            arquivo.seek(0)
            doc = DocxDocument(arquivo)
            
            texto_completo = ""
            for paragrafo in doc.paragraphs:
                texto_completo += paragrafo.text + "\n"
            
            # Adicionar tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        texto_completo += celula.text + " "
                    texto_completo += "\n"
            
            if texto_completo.strip():
                doc_langchain = Document(
                    page_content=texto_completo,
                    metadata={'source': arquivo.name, 'tipo': 'Word'}
                )
                return [doc_langchain]
        
        except ImportError:
            # Método 2: Unstructured
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                arquivo.seek(0)
                tmp_file.write(arquivo.read())
                tmp_file_path = tmp_file.name
            
            from langchain_community.document_loaders import UnstructuredWordDocumentLoader
            loader = UnstructuredWordDocumentLoader(tmp_file_path)
            docs = loader.load()
            
            for doc in docs:
                doc.metadata['source'] = arquivo.name
                doc.metadata['tipo'] = 'Word'
            
            os.unlink(tmp_file_path)
            return docs
        
        return []
        
    except Exception as e:
        st.error(f"Erro ao processar Word {arquivo.name}: {str(e)}")
        return []

def processar_excel(arquivo):
    """Processa arquivo Excel"""
    try:
        arquivo.seek(0)
        
        # Tentar ler todas as planilhas
        excel_file = pd.ExcelFile(arquivo)
        texto_completo = f"ARQUIVO EXCEL: {arquivo.name}\n\n"
        
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(arquivo, sheet_name=sheet_name)
                texto_completo += f"=== PLANILHA: {sheet_name} ===\n"
                texto_completo += df.to_string(index=False) + "\n\n"
                
                # Adicionar informações estatísticas
                texto_completo += f"Linhas: {len(df)}, Colunas: {len(df.columns)}\n"
                texto_completo += f"Colunas: {', '.join(df.columns.astype(str))}\n\n"
                
            except Exception as e:
                texto_completo += f"Erro ao ler planilha {sheet_name}: {str(e)}\n\n"
        
        doc = Document(
            page_content=texto_completo,
            metadata={'source': arquivo.name, 'tipo': 'Excel'}
        )
        
        return [doc]
        
    except Exception as e:
        st.error(f"Erro ao processar Excel {arquivo.name}: {str(e)}")
        return []

def criar_chat_com_documentos(documentos):
    """Cria o chat com os documentos processados"""
    try:
        # Dividir documentos em chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
        docs_split = text_splitter.split_documents(documentos)
        
        # Criar vector store
        embeddings = OpenAIEmbeddings()
        vector_store = FAISS.from_documents(docs_split, embeddings)
        
        # Configurar LLM
        llm = ChatOpenAI(
            model="gpt-3.5-turbo",
            temperature=0.0  # Zero criatividade
        )
        
        # Prompt customizado
        prompt_template = """
Você é um assistente que responde perguntas baseado EXCLUSIVAMENTE nos documentos fornecidos.

REGRAS IMPORTANTES:
1. Use APENAS as informações dos documentos fornecidos
2. Se a informação não estiver nos documentos, diga "Esta informação não está disponível nos documentos fornecidos"
3. NUNCA use conhecimento externo ou invente informações
4. Cite sempre os arquivos de onde veio a informação
5. Seja preciso e específico

DOCUMENTOS:
{context}

PERGUNTA: {question}

RESPOSTA (baseada APENAS nos documentos fornecidos):
"""
        
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        # Configurar retriever
        retriever = vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 8}
        )
        
        # Criar chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=True,
            chain_type_kwargs={"prompt": prompt}
        )
        
        st.session_state["chat_chain_upload"] = qa_chain
        
    except Exception as e:
        st.error(f"Erro ao criar chat: {str(e)}")
